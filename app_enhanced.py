import streamlit as st
import pandas as pd
from docx import Document
from docx.table import Table
from docx.shared import Pt, RGBColor
import io
import re


# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

ORANGE = RGBColor(230, 81, 0)
GREEN  = RGBColor(0, 135, 81)

HIGH_ATTENDANCE_PCT     = 0.80  # attended >= 80% of session → High
MODERATE_ATTENDANCE_PCT = 0.65  # attended >= 65% and < 80%  → Moderate
# attended > 0% and < 65%  → Low

SPLIT_THRESHOLD = 17  # lists with >= 17 names use both left and right columns

# Known Meet display name → official tracker name mappings.
# Add a new entry whenever a mentee joins with an unrecognised display name.
NAME_FIXES = {
    "vickysmane":               "victory chiamaka eze",
    "vickysmane chinedu":       "victory chiamaka eze",
    "dd sylvia":                "chinedu divinefavour",
    "chukky igunbor":           "georgina samuel",
    "chukky igbunbor":          "georgina samuel",
    "pearl owusu -twum":        "pearl owusu-twum",
    "pearl owusu  twum":        "pearl owusu-twum",
    "damilola adesodun":        "damilola bankole adesodun",
    "sharon nngwama":           "sharon ngwama",
    "tega ishaya":              "phoebe ishaya oghenetega",
    "phoebe ishaya oghenetega": "phoebe ishaya oghenetega",
    "moyin ojo":                "moyinoluwa ojo",
    "funke":                    "mary oloyede funke",
    "lela":                     "lela tony",
    # "divine omeire" intentionally omitted — different person from chinedu divine favour
}

APP_CSS = """
<style>
.stApp { background: linear-gradient(to bottom, #ffffff, #f0f2f6); }
.section-header { color: #008751; font-size: 1.2rem; font-weight: 700; text-transform: uppercase; margin-bottom: 15px; }
.stButton>button { background: linear-gradient(135deg, #E65100 0%, #FF9800 100%); color: white; border: none; padding: 12px 30px; font-weight: bold; border-radius: 50px; }
[data-testid="stMetric"] { background-color: white; padding: 15px; border-radius: 10px; border-left: 5px solid #E65100; box-shadow: 0 2px 5px rgba(0,0,0,0.05); }
[data-testid="stWidgetLabel"] > label, [data-testid="stWidgetLabel"] { color: #1a1a1a !important; }
</style>
"""


# ─────────────────────────────────────────────────────────────────────────────
# DATA HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def duration_to_seconds(duration_str: str) -> int:
    if pd.isna(duration_str) or duration_str == "":
        return 0
    text = str(duration_str)
    hours   = re.search(r'(\d+)\s*(?:hours?|hrs?)',    text, re.IGNORECASE)
    minutes = re.search(r'(\d+)\s*(?:minutes?|mins?)', text, re.IGNORECASE)
    seconds = re.search(r'(\d+)\s*(?:seconds?|secs?|s\b)', text, re.IGNORECASE)
    total = 0
    if hours:   total += int(hours.group(1))   * 3600
    if minutes: total += int(minutes.group(1)) * 60
    if seconds: total += int(seconds.group(1))

    return total


def extract_full_name(row: pd.Series) -> str:
    first = str(row['First Name']).strip() if 'First Name' in row else ""
    last  = str(row['Last Name']).strip()  if 'Last Name'  in row else ""
    return f"{first} {last}".strip()


def load_tracker(tracker_file) -> tuple[set, dict]:
    tracker = pd.read_csv(tracker_file)
    tracker.columns = tracker.columns.str.strip()

    missing_columns = [c for c in ['First Name', 'Last Name'] if c not in tracker.columns]
    if missing_columns:
        st.error(f"Tracker CSV is missing columns: {missing_columns}. Found: {list(tracker.columns)}")
        st.stop()

    tracker = tracker.dropna(subset=['First Name', 'Last Name'])
    tracker['Full Name'] = (
        tracker['First Name'].str.strip() + " " + tracker['Last Name'].str.strip()
    ).str.lower()

    mentee_list = set(tracker['Full Name'])
    pod_map     = tracker.set_index('Full Name')['Pod'].to_dict() if 'Pod' in tracker.columns else {}
    return mentee_list, pod_map


def resolve_meet_name(raw_name: str, mentee_list: set, collapsed_map: dict) -> str:
    if raw_name in NAME_FIXES:
        return NAME_FIXES[raw_name]
    if raw_name in mentee_list:
        return raw_name
    # Handles compound surnames where spaces were joined or split differently
    if raw_name.replace(' ', '') in collapsed_map:
        return collapsed_map[raw_name.replace(' ', '')]

    raw_parts = set(raw_name.replace('-', ' ').split())
    raw_last  = raw_name.split()[-1] if raw_name.split() else ""
    best_match, best_score = None, 0

    for mentee in mentee_list:
        mentee_parts = set(mentee.replace('-', ' ').split())
        mentee_last  = mentee.split()[-1] if mentee.split() else ""
        overlap      = len(raw_parts & mentee_parts)
        score        = overlap + (2 if raw_last == mentee_last else 0)
        if score > best_score and (raw_last == mentee_last or overlap >= 2):
            best_score = score
            best_match = mentee

    return best_match if best_match else raw_name


def is_team_member(name: str, team_set: set) -> bool:
    if name in team_set:
        return True
    name_parts = set(name.replace('-', ' ').split())
    name_last  = name.split()[-1] if name.split() else ""
    for team_name in team_set:
        team_parts = set(team_name.replace('-', ' ').split())
        team_last  = team_name.split()[-1] if team_name.split() else ""
        overlap    = len(name_parts & team_parts)
        if team_last == name_last or overlap >= 2:
            return True
    return False


def classify_attendance(attendance_df, mentee_list: set, team_set: set, excluded_names: set):
    mentees_found = attendance_df[attendance_df['Name'].isin(mentee_list)]
    unmatched     = attendance_df[~attendance_df['Name'].isin(excluded_names)].copy()
    facilitators  = unmatched[unmatched['Name'].apply(lambda n: is_team_member(n, team_set))]
    guests        = unmatched[~unmatched['Name'].apply(lambda n: is_team_member(n, team_set))]
    return mentees_found, facilitators, guests


def group_names_by_pod(name_list: list, pod_map: dict) -> tuple[str, str]:
    pods_dict = {}
    for name in name_list:
        pod = pod_map.get(name.lower(), "Unknown Pod")
        pods_dict.setdefault(pod, []).append(name)
    pods_text  = "\n".join(str(pod) for pod in pods_dict)
    names_text = "\n".join(", ".join(members) for members in pods_dict.values())
    return pods_text, names_text


def build_attendance_results(meet_log, mentees_found, mentee_list, pod_map, team_list, meeting_duration_seconds: int) -> dict:
    high_threshold     = HIGH_ATTENDANCE_PCT     * meeting_duration_seconds
    moderate_threshold = MODERATE_ATTENDANCE_PCT * meeting_duration_seconds
    seconds            = mentees_found['Seconds']
    return {
        "total_participants":   len(meet_log['RawName'].unique()),
        "expected_mentees":     len(mentee_list),
        "present_mentees":      len(mentees_found),
        "attendance_rate":      f"{round((len(mentees_found) / len(mentee_list)) * 100, 1)}%",
        "high_attendance":      sorted(mentees_found[seconds >= high_threshold]['Name'].str.title().tolist()),
        "moderate_attendance":  sorted(mentees_found[(seconds >= moderate_threshold) & (seconds < high_threshold)]['Name'].str.title().tolist()),
        "low_attendance":       sorted(mentees_found[seconds < moderate_threshold]['Name'].str.title().tolist()),
        "absent":               sorted([name.title() for name in mentee_list if name not in set(mentees_found['Name'])]),
        "pod_map":              pod_map,
        "team":                 team_list,
    }


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def consolidate_and_replace(paragraphs, replacements: dict):
    for paragraph in paragraphs:
        # Word splits placeholder strings across multiple runs; merge them first
        if paragraph.runs and any(key in paragraph.text for key in replacements):
            paragraph.runs[0].text = "".join(run.text for run in paragraph.runs)
            for run in paragraph.runs[1:]:
                run.text = ""
        for key, value in replacements.items():
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))


def iter_body_tables(body_element):
    for child in body_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            yield Table(child, body_element)
        else:
            yield from iter_body_tables(child)


def fill_list_left_cell(cell, names: list, color: RGBColor = None):
    cell.text    = ""
    use_two_cols = len(names) >= SPLIT_THRESHOLD
    midpoint     = (len(names) + 1) // 2
    display      = names[:midpoint] if use_two_cols else names
    font_size    = Pt(9) if use_two_cols else Pt(10)
    for i, name in enumerate(display):
        line      = f"{i + 1}. {name}"
        paragraph = cell.add_paragraph(line)
        run       = paragraph.runs[0] if paragraph.runs else paragraph.add_run(line)
        run.font.name = 'Montserrat'
        run.font.size = font_size
        if color:
            run.font.color.rgb = color


def fill_list_right_cell(cell, names: list, color: RGBColor = None):
    cell.text = ""
    if len(names) < SPLIT_THRESHOLD:
        return
    midpoint    = (len(names) + 1) // 2
    right_names = names[midpoint:]
    for i, name in enumerate(right_names):
        line      = f"{midpoint + i + 1}. {name}"
        paragraph = cell.add_paragraph(line)
        run       = paragraph.runs[0] if paragraph.runs else paragraph.add_run(line)
        run.font.name = 'Montserrat'
        run.font.size = Pt(9)
        if color:
            run.font.color.rgb = color


def apply_stat_highlight(cell, replacements: dict):
    stat_keys = ["{{total_participants}}", "{{total_mentees}}", "{{total_present}}", "{{attendance_rate}}"]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            for key in stat_keys:
                if run.text == replacements.get(key, "___NO___"):
                    run.font.bold      = True
                    run.font.color.rgb = ORANGE
                    run.font.size      = Pt(18)


def build_replacements(session: dict, results: dict, observations: dict) -> dict:
    low_pods, _  = group_names_by_pod(results['low_attendance'], results['pod_map'])
    abs_pods, _  = group_names_by_pod(results['absent'],         results['pod_map'])
    session_date_only = session['date'].split(',')[-1].strip() if ',' in session['date'] else session['date']
    return {
        "{{week_number}}":        session['week_number'],
        "{{total_weeks}}":        session['total_weeks'],
        "{{session_name}}":       session['topic'],
        "{{session_day_date}}":   session['date'],
        "{{session_time}}":       session['time_range'],
        "{{total_participants}}": str(results['total_participants']),
        "{{total_mentees}}":      str(results['expected_mentees']),
        "{{total_present}}":      str(results['present_mentees']),
        "{{attendance_rate}}":    results['attendance_rate'],
        "{{facilitator}}":        session['facilitator'],
        "{{session_date}}":       session_date_only,
        "{{session_duration}}":   f"{session['duration']} ({session['time_range']})",
        "{{high_count}}":         str(len(results['high_attendance'])),
        "{{moderate_count}}":     str(len(results['moderate_attendance'])),
        "{{low_count}}":          str(len(results['low_attendance'])),
        "{{low_pod}}":            low_pods,
        "{{absent_pod}}":         abs_pods,
        "{{obs_attendance}}":     observations['attendance'],
        "{{obs_engagement}}":     observations['engagement'],
        "{{obs_absentees}}":      observations['absentees'],
    }


def generate_report(template_file, replacements: dict, results: dict) -> bytes:
    doc = Document(template_file)

    # Replace placeholders in document headers
    for section in doc.sections:
        consolidate_and_replace(section.header.paragraphs, replacements)
        for table in section.header.tables:
            for row in table.rows:
                seen_cells = set()
                for cell in row.cells:
                    if id(cell._tc) in seen_cells:
                        continue
                    seen_cells.add(id(cell._tc))
                    consolidate_and_replace(cell.paragraphs, replacements)

    # Replace placeholders in body paragraphs
    consolidate_and_replace(doc.paragraphs, replacements)

    # Maps each template tag to (side, name_list, optional_color).
    # left cells: all names when < SPLIT_THRESHOLD, first half when >= SPLIT_THRESHOLD.
    # right cells: cleared when < SPLIT_THRESHOLD, second half when >= SPLIT_THRESHOLD.
    list_cell_map = {
        "{{team_members_left}}":   ('left',  results['team'],                GREEN),
        "{{team_members_right}}":  ('right', results['team'],                GREEN),
        "{{high_list_left}}":      ('left',  results['high_attendance'],      None),
        "{{high_list_right}}":     ('right', results['high_attendance'],      None),
        "{{moderate_list_left}}":  ('left',  results['moderate_attendance'],  None),
        "{{moderate_list_right}}": ('right', results['moderate_attendance'],  None),
        "{{low_list_left}}":       ('left',  results['low_attendance'],       None),
        "{{low_list_right}}":      ('right', results['low_attendance'],       None),
        "{{absent_list_left}}":    ('left',  results['absent'],               None),
        "{{absent_list_right}}":   ('right', results['absent'],               None),
    }

    # Replace placeholders in all body tables
    for table in iter_body_tables(doc.element.body):
        for row in table.rows:
            seen_cells = set()
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                cell_text = cell.text

                list_tag_handled = False
                for tag, (side, names, color) in list_cell_map.items():
                    if tag in cell_text:
                        if side == 'left':
                            fill_list_left_cell(cell, names, color)
                        else:
                            fill_list_right_cell(cell, names, color)
                        list_tag_handled = True
                        break

                if not list_tag_handled:
                    consolidate_and_replace(cell.paragraphs, replacements)
                    apply_stat_highlight(cell, replacements)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# UI SECTIONS
# ─────────────────────────────────────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.image("logooooo-removebg-preview.png", use_container_width=True)
        st.divider()
        st.caption("Admin: Stephanie Nwangwu")


def render_session_identity() -> dict:
    st.markdown("<p class='section-header'>Step 1: Session Identity</p>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        week_number      = st.text_input("Week Number",    value="2")
        total_weeks      = st.text_input("Total Weeks",    value="4")
        topic            = st.text_input("Session Topic",  value="Execution Framework")
    with col2:
        session_date     = st.text_input("Day & Date",         value="Sunday, 12th April 2026")
        time_range       = st.text_input("Session Time Range", value="8:00 PM – 11:10 PM")
    with col3:
        facilitator      = st.text_input("Facilitator",    value="Temilade Salami")
        session_duration = st.text_input("Total Duration", value="3 hours 10 minutes")

    team_input = st.text_area(
        "Team Member List (comma separated)",
        value="Sybil Obeng-Sintim, Paseal Njoku, Janet Isesele, Oluwasanmi Awe, Vivian Nesiama, Edidiong Udoudom, Grace Adu-Yeboah, Stephanie Nwangwu"
    )

    return {
        "week_number": week_number,
        "total_weeks": total_weeks,
        "topic":       topic,
        "date":        session_date,
        "time_range":  time_range,
        "facilitator": facilitator,
        "duration":    session_duration,
        "team_input":  team_input,
    }


def render_file_uploaders() -> tuple:
    st.markdown("<p class='section-header'>Step 2: Data Ingestion</p>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        tracker_file  = st.file_uploader("Official Mentee Tracker (CSV)",  type="csv")
    with col2:
        meet_log_file = st.file_uploader("Google Meet Log (CSV)",           type="csv")
    with col3:
        template_file = st.file_uploader("Word Master Template (.docx)",    type="docx")
    return tracker_file, meet_log_file, template_file


def render_match_feedback(facilitators, guests):
    if len(facilitators) == 0 and len(guests) == 0:
        st.success("✅ All Meet names matched to mentees successfully")
        return
    if len(facilitators) > 0:
        st.info(f"ℹ️ {len(facilitators)} team member(s) detected in the Meet log:")
        st.dataframe(facilitators[['Name']].rename(columns={'Name': 'Team Member'}), hide_index=True)
    if len(guests) > 0:
        st.warning(f"⚠️ {len(guests)} guest(s) detected — not a mentee or team member. Add to NAME_FIXES if they are real mentees:")
        st.dataframe(guests[['Name']].rename(columns={'Name': 'Guest'}), hide_index=True)


def render_dashboard(results: dict):
    st.divider()
    st.markdown("<p class='section-header'>Sprint Dashboard</p>", unsafe_allow_html=True)
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Participants", results['total_participants'])
    col2.metric("Expected Mentees",   results['expected_mentees'])
    col3.metric("Mentees Present",    results['present_mentees'])
    col4.metric("Attendance Rate",    results['attendance_rate'])


def render_observations(results: dict) -> dict:
    st.markdown("<p class='section-header'>Step 3: Operations Observations</p>", unsafe_allow_html=True)
    absent_count   = len(results['absent'])
    attendance_obs = st.text_area(
        "Attendance Observation",
        value=f"{results['present_mentees']}/{results['expected_mentees']} mentees present."
    )
    engagement_obs = st.text_area(
        "Engagement Observation",
        value="High; majority stayed for the duration."
    )
    absentees_obs  = st.text_area(
        "Absentees Note",
        value=f"{absent_count} mentee{'s' if absent_count != 1 else ''} missed the session."
    )
    return {
        "attendance": attendance_obs,
        "engagement": engagement_obs,
        "absentees":  absentees_obs,
    }


# ─────────────────────────────────────────────────────────────────────────────
# APP ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="Launchpad Sprint Operations", layout="wide", page_icon="🚀")
st.markdown(APP_CSS, unsafe_allow_html=True)

render_sidebar()
st.markdown("<h1 style='text-align: center; color: #E65100;'>Launchpad Sprint Operations</h1>", unsafe_allow_html=True)
st.divider()

session                                    = render_session_identity()
tracker_file, meet_log_file, template_file = render_file_uploaders()

if 'processed' not in st.session_state:
    st.session_state.processed = False

if st.button("🚀 PROCESS SESSION DATA"):
    if not all([tracker_file, meet_log_file, template_file]):
        st.error("Please upload all three files before processing.")
    else:
        team_list = [name.strip() for name in session['team_input'].split(",")]

        mentee_list, pod_map  = load_tracker(tracker_file)
        collapsed_map         = {name.replace(' ', ''): name for name in mentee_list}

        print("Collapsed map: ", collapsed_map)

        meet_log              = pd.read_csv(meet_log_file).dropna(subset=['Participant Name'])
        meet_log['RawName']   = meet_log['Participant Name'].str.strip().str.lower()
        meet_log['Name']      = meet_log['RawName'].apply(lambda n: resolve_meet_name(n, mentee_list, collapsed_map))
        meet_log['Seconds']   = meet_log['Attended Duration'].apply(duration_to_seconds)

        attendance      = meet_log.groupby('Name')['Seconds'].sum().reset_index()
        team_set        = {name.lower() for name in team_list}
        excluded_names  = mentee_list | {session['facilitator'].strip().lower()}

        mentees_found, facilitators, guests = classify_attendance(
            attendance, mentee_list, team_set, excluded_names
        )
        render_match_feedback(facilitators, guests)

        meeting_duration_seconds   = duration_to_seconds(session['duration'])
        st.session_state.results   = build_attendance_results(meet_log, mentees_found, mentee_list, pod_map, team_list, meeting_duration_seconds)
        st.session_state.processed = True

if st.session_state.processed:
    results = st.session_state.results
    render_dashboard(results)
    observations = render_observations(results)

    if st.button("📝 GENERATE SPRINT REPORT"):
        replacements = build_replacements(session, results, observations)
        report_bytes = generate_report(template_file, replacements, results)
        st.balloons()
        st.download_button(
            "📥 DOWNLOAD FINAL REPORT",
            data=report_bytes,
            file_name=f"Sprint_Week_{session['week_number']}_Report.docx"
        )
